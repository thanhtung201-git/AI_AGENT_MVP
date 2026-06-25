"""
word_reader.py — Đọc file PO dạng Word (.docx)
Dùng python-docx để extract paragraph và bảng biểu

Cài đặt: pip install python-docx
"""

from docx import Document
import logging

logger = logging.getLogger(__name__)


def read_word(file_path: str) -> dict:
    """
    Đọc toàn bộ nội dung file Word PO.
    Extract cả đoạn văn (paragraph) và bảng biểu (table).
    Giữ đúng thứ tự xuất hiện trong file.

    Args:
        file_path (str): Đường dẫn file Word (.docx)

    Returns:
        dict: {
            "success":     True/False,
            "text":        Toàn bộ nội dung text,
            "format":      "word",
            "paragraphs":  Số đoạn văn đọc được,
            "tables":      Số bảng đọc được,
            "error":       Lỗi nếu có
        }
    """
    try:
        doc = Document(file_path)
        logger.info(f"Đang đọc Word: {file_path}")

        full_text = []
        para_count = 0
        table_count = 0

        # ── Duyệt theo đúng thứ tự trong file ──
        # (paragraphs và tables xen kẽ nhau trong Word)
        for block in _iter_blocks(doc):

            if block["type"] == "paragraph":
                text = block["text"].strip()
                if text:  # Bỏ dòng trống
                    full_text.append(text)
                    para_count += 1

            elif block["type"] == "table":
                table_count += 1
                full_text.append(f"[Bảng {table_count}]")

                for row in block["rows"]:
                    clean_row = [cell.strip() for cell in row]
                    full_text.append("\t".join(clean_row))

                full_text.append("")  # Dòng trống sau bảng

        result_text = "\n".join(full_text).strip()

        if not result_text:
            return {
                "success": False,
                "text": "",
                "format": "word",
                "paragraphs": 0,
                "tables": 0,
                "error": "File Word không có nội dung"
            }

        return {
            "success": True,
            "text": result_text,
            "format": "word",
            "paragraphs": para_count,
            "tables": table_count,
            "error": None
        }

    except FileNotFoundError:
        return {
            "success": False,
            "text": "",
            "format": "word",
            "paragraphs": 0,
            "tables": 0,
            "error": f"Không tìm thấy file: {file_path}"
        }

    except Exception as e:
        logger.error(f"Lỗi đọc Word {file_path}: {e}")
        return {
            "success": False,
            "text": "",
            "format": "word",
            "paragraphs": 0,
            "tables": 0,
            "error": f"Lỗi khi đọc Word: {str(e)}"
        }


def _iter_blocks(doc):
    """
    Generator duyệt qua các block (paragraph + table) theo đúng thứ tự
    trong file Word — vì python-docx không tự duy trì thứ tự khi dùng
    doc.paragraphs và doc.tables riêng biệt.
    """
    from docx.oxml.ns import qn

    body = doc.element.body
    for child in body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

        if tag == "p":  # Paragraph
            text = "".join(r.text for r in child.iter(qn("w:t")))
            yield {"type": "paragraph", "text": text}

        elif tag == "tbl":  # Table
            rows = []
            for row_el in child.iter(qn("w:tr")):
                cells = []
                for cell_el in row_el.iter(qn("w:tc")):
                    cell_text = "".join(
                        r.text for r in cell_el.iter(qn("w:t"))
                    )
                    cells.append(cell_text)
                if cells:
                    rows.append(cells)
            yield {"type": "table", "rows": rows}


# ── Chạy thử ──
if __name__ == "__main__":
    result = read_word("sample_data/test_po.docx")
    print(f"✅ Paragraphs: {result['paragraphs']} | Tables: {result['tables']}")
    print(result["text"][:500])